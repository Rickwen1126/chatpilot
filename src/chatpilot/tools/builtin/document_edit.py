"""document_edit tool — edit .docx / .xlsx files and upload to R2."""

from __future__ import annotations

import io
import json
import logging
from pathlib import Path
from typing import Any

from copilot.types import ToolInvocation, ToolResult

from chatpilot.core.types import AccessLevel, ToolDefinition
from chatpilot.files.center import FileHandleCenter
from chatpilot.files.models import FileKind
from chatpilot.tools.session_context import get_session_context

logger = logging.getLogger(__name__)

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB


def _parse_ref(ref: str) -> tuple[str, str, str]:
    """Parse file_ref → (platform, message_id, filename).

    Format: "platform:message_id:filename"
    """
    parts = ref.split(":", 2)
    if ref.startswith("line:") and ref.count(":") >= 3:
        parts = ref.split(":", 3)
        platform = ":".join(parts[:2])
        message_id = parts[2]
        filename = parts[3] if len(parts) == 4 else ""
        return platform, message_id, filename
    if len(parts) < 2:
        raise ValueError(f"無效的 file_ref 格式: {ref}（應為 platform:message_id:filename）")
    platform = parts[0]
    message_id = parts[1]
    filename = parts[2] if len(parts) == 3 else ""
    return platform, message_id, filename


def _get_extension(filename: str, data: bytes) -> str:
    """Determine file extension from filename or magic bytes."""
    if filename:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext in ("xlsx", "docx"):
            return ext
    # Fallback: check ZIP magic + content hints
    if data[:4] == b"PK\x03\x04":
        if b"word/" in data[:2000]:
            return "docx"
        if b"xl/" in data[:2000]:
            return "xlsx"
    return ""


async def _edit_xlsx(data: bytes, instruction: str, edit_data: str | None) -> bytes:
    """Edit .xlsx file using openpyxl."""
    import asyncio

    def _do_edit() -> bytes:
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(data))
        ws = wb.active

        if edit_data:
            parsed = json.loads(edit_data)
            if isinstance(parsed, list):
                # Append rows
                for row in parsed:
                    if isinstance(row, list):
                        ws.append(row)
                    elif isinstance(row, dict):
                        ws.append(list(row.values()))
            elif isinstance(parsed, dict):
                # Update specific cells
                for cell_ref, value in parsed.items():
                    ws[cell_ref] = value
        else:
            # No structured data — append instruction as a note row
            ws.append([instruction])

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    return await asyncio.to_thread(_do_edit)


async def _edit_docx(data: bytes, instruction: str, edit_data: str | None) -> bytes:
    """Edit .docx file using python-docx."""
    import asyncio

    def _do_edit() -> bytes:
        import docx

        doc = docx.Document(io.BytesIO(data))

        if edit_data:
            parsed = json.loads(edit_data)
            if isinstance(parsed, str):
                doc.add_paragraph(parsed)
            elif isinstance(parsed, list):
                for item in parsed:
                    doc.add_paragraph(str(item))
            elif isinstance(parsed, dict):
                # Replace text
                old = parsed.get("old", "")
                new = parsed.get("new", "")
                if old and new:
                    for para in doc.paragraphs:
                        if old in para.text:
                            for run in para.runs:
                                if old in run.text:
                                    run.text = run.text.replace(old, new)
        else:
            doc.add_paragraph(instruction)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    return await asyncio.to_thread(_do_edit)


def create_document_edit_tool(
    adapters: dict,
    r2_storage: Any,
    file_handle_center: FileHandleCenter | None = None,
) -> ToolDefinition:
    """Create document_edit tool for editing uploaded files."""

    async def handler(invocation: ToolInvocation) -> ToolResult:
        args = invocation.get("arguments") or {}
        file_ref = args.get("file_ref", "")
        instruction = args.get("instruction", "")
        edit_data = args.get("data")

        if not file_ref:
            return ToolResult(
                textResultForLlm="缺少 file_ref 參數",
                resultType="failure",
            )
        if not instruction:
            return ToolResult(
                textResultForLlm="缺少 instruction 參數",
                resultType="failure",
            )

        try:
            platform, message_id, filename = _parse_ref(file_ref)
        except ValueError as e:
            return ToolResult(textResultForLlm=str(e), resultType="failure")

        file_bytes: bytes | None = None
        session_context = None
        source_handle = None
        if file_handle_center is not None:
            try:
                session_context = get_session_context(invocation)
            except RuntimeError:
                session_context = None
            if session_context is not None:
                source_handle = await file_handle_center.find_source_handle(
                    route_id=session_context.route_id,
                    platform=platform,
                    native_locator=message_id,
                )
                if source_handle is not None:
                    local_path = await file_handle_center.ensure_local(
                        source_handle.file_id
                    )
                    file_bytes = Path(local_path).read_bytes()
                    filename = filename or source_handle.filename or ""

        if file_bytes is None:
            adapter = adapters.get(platform)
            if adapter is None:
                return ToolResult(
                    textResultForLlm=f"未知的平台: {platform}",
                    resultType="failure",
                )
            logger.warning(
                "[file] legacy document_edit fallback file_ref=%s route=%s",
                file_ref,
                session_context.route_id if session_context is not None else "unknown",
            )
            file_bytes = await adapter.download_media(message_id)
        if file_bytes is None:
            return ToolResult(
                textResultForLlm="無法下載檔案，可能已過期，請重新上傳",
                resultType="failure",
            )

        if len(file_bytes) > MAX_FILE_SIZE:
            return ToolResult(
                textResultForLlm="檔案過大（超過 10MB），請縮小檔案後再試",
                resultType="failure",
            )

        # Determine format
        ext = _get_extension(filename, file_bytes)
        if ext not in ("xlsx", "docx"):
            return ToolResult(
                textResultForLlm="目前只支援 .docx 和 .xlsx 格式",
                resultType="failure",
            )

        # Edit
        try:
            if ext == "xlsx":
                output = await _edit_xlsx(file_bytes, instruction, edit_data)
            else:
                output = await _edit_docx(file_bytes, instruction, edit_data)
        except json.JSONDecodeError:
            return ToolResult(
                textResultForLlm="data 參數 JSON 格式不正確",
                resultType="failure",
            )
        except Exception as e:
            logger.error("Document edit failed: %s", e)
            return ToolResult(
                textResultForLlm=f"檔案格式異常，無法開啟: {e}",
                resultType="failure",
            )

        # Upload to R2
        content_types = {
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
        out_name = filename or f"edited.{ext}"
        generated_handle = None
        if file_handle_center is not None and session_context is not None:
            generated_handle = await file_handle_center.register_bytes_file(
                route_id=session_context.route_id,
                data=output,
                kind=FileKind.file,
                filename=out_name,
                mime_type=content_types[ext],
                derived_from_file_id=(
                    source_handle.file_id if source_handle is not None else None
                ),
                generated_by_tool="document_edit",
            )
        url = await r2_storage.upload(
            data=output,
            content_type=content_types[ext],
            extension=ext,
        )
        if url is None:
            return ToolResult(
                textResultForLlm=(
                    "檔案處理完成但上傳失敗，請稍後再試"
                    + (
                        f"\n已納管 file_id: {generated_handle.file_id}"
                        if generated_handle is not None
                        else ""
                    )
                ),
                resultType="failure",
            )

        return ToolResult(
            textResultForLlm=(
                f"已完成編輯並上傳: {out_name}\n"
                f"下載連結: {url}\n"
                f"（連結 7 天內有效）"
                + (
                    f"\nfile_id: {generated_handle.file_id}"
                    if generated_handle is not None
                    else ""
                )
            ),
            resultType="success",
        )

    return ToolDefinition(
        name="document_edit",
        description=(
            "編輯 .docx 或 .xlsx 檔案。"
            "接收檔案參考 ref 和編輯指令，處理後上傳並回傳下載連結。"
            "支援操作：追加工作表列、新增段落、修改內容等。"
        ),
        parameters={
            "type": "object",
            "properties": {
                "file_ref": {
                    "type": "string",
                    "description": "檔案參考 ID（如 line:msg_123:report.xlsx）",
                },
                "instruction": {
                    "type": "string",
                    "description": "編輯指令（如：在最後加上總結段落、追加一列出貨資料）",
                },
                "data": {
                    "type": "string",
                    "description": "需要寫入的具體資料（JSON 格式），如追加的表格列內容",
                },
            },
            "required": ["file_ref", "instruction"],
        },
        handler=handler,
        access_level=AccessLevel.GLOBAL,
    )
