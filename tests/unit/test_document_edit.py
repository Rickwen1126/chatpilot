"""Tests for document_edit tool — xlsx/docx editing + full handler flow."""

from __future__ import annotations

import io
import json

import docx
import openpyxl
import pytest

from chatpilot.files.center import FileHandleCenter
from chatpilot.files.models import FileKind, SourceFetchResult, SourceHandleInput
from chatpilot.files.store import SqliteFileStore
from chatpilot.tools.builtin.document_edit import (
    _edit_docx,
    _edit_xlsx,
    _get_extension,
    _parse_ref,
    create_document_edit_tool,
)

# ── Helpers ──────────────────────────────────────────────────────


def _make_xlsx(rows: list[list]) -> bytes:
    """Create a minimal xlsx file in memory."""
    wb = openpyxl.Workbook()
    ws = wb.active
    for row in rows:
        ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _make_docx(paragraphs: list[str]) -> bytes:
    """Create a minimal docx file in memory."""
    doc = docx.Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _read_xlsx_rows(data: bytes) -> list[list]:
    wb = openpyxl.load_workbook(io.BytesIO(data))
    ws = wb.active
    return [[cell.value for cell in row] for row in ws.iter_rows()]


def _read_docx_paragraphs(data: bytes) -> list[str]:
    doc = docx.Document(io.BytesIO(data))
    return [p.text for p in doc.paragraphs]


# ── _parse_ref ───────────────────────────────────────────────────


def test_parse_ref_three_parts():
    assert _parse_ref("line:msg_123:report.xlsx") == ("line", "msg_123", "report.xlsx")


def test_parse_ref_two_parts():
    assert _parse_ref("line:msg_456") == ("line", "msg_456", "")


def test_parse_ref_invalid():
    with pytest.raises(ValueError, match="無效"):
        _parse_ref("bad")


def test_parse_ref_named_line_channel():
    assert _parse_ref("line:webric:msg_123:report.xlsx") == (
        "line:webric",
        "msg_123",
        "report.xlsx",
    )


# ── _get_extension ───────────────────────────────────────────────


def test_get_extension_from_filename():
    assert _get_extension("report.xlsx", b"") == "xlsx"
    assert _get_extension("doc.docx", b"") == "docx"


def test_get_extension_unknown_filename():
    assert _get_extension("file.pdf", b"") == ""


def test_get_extension_from_magic_bytes():
    xlsx_data = _make_xlsx([["test"]])
    assert _get_extension("", xlsx_data) == "xlsx"

    docx_data = _make_docx(["test"])
    assert _get_extension("", docx_data) == "docx"


# ── _edit_xlsx ───────────────────────────────────────────────────


async def test_xlsx_append_rows():
    original = _make_xlsx([["品名", "數量"], ["水泥漆", 10]])
    data_json = json.dumps([["乳膠漆", 5], ["防水漆", 3]])

    result = await _edit_xlsx(original, "追加", data_json)
    rows = _read_xlsx_rows(result)

    assert len(rows) == 4
    assert rows[2] == ["乳膠漆", 5]
    assert rows[3] == ["防水漆", 3]


async def test_xlsx_append_dict_rows():
    original = _make_xlsx([["品名", "數量"]])
    data_json = json.dumps([{"品名": "虹牌水泥漆", "數量": 20}])

    result = await _edit_xlsx(original, "追加", data_json)
    rows = _read_xlsx_rows(result)

    assert len(rows) == 2
    assert rows[1] == ["虹牌水泥漆", 20]


async def test_xlsx_update_cells():
    original = _make_xlsx([["品名", "數量"], ["水泥漆", 10]])
    data_json = json.dumps({"B2": 99})

    result = await _edit_xlsx(original, "更新", data_json)
    rows = _read_xlsx_rows(result)

    assert rows[1][1] == 99


async def test_xlsx_no_data_appends_instruction():
    original = _make_xlsx([["A"]])

    result = await _edit_xlsx(original, "這是備註", None)
    rows = _read_xlsx_rows(result)

    assert rows[1] == ["這是備註"]


# ── _edit_docx ───────────────────────────────────────────────────


async def test_docx_append_string():
    original = _make_docx(["第一段"])
    data_json = json.dumps("總結：完成")

    result = await _edit_docx(original, "加總結", data_json)
    paragraphs = _read_docx_paragraphs(result)

    assert "總結：完成" in paragraphs


async def test_docx_append_list():
    original = _make_docx(["第一段"])
    data_json = json.dumps(["第二段", "第三段"])

    result = await _edit_docx(original, "追加", data_json)
    paragraphs = _read_docx_paragraphs(result)

    assert "第二段" in paragraphs
    assert "第三段" in paragraphs


async def test_docx_replace_text():
    original = _make_docx(["測試報告初稿"])
    data_json = json.dumps({"old": "初稿", "new": "定稿"})

    result = await _edit_docx(original, "取代", data_json)
    paragraphs = _read_docx_paragraphs(result)

    assert any("定稿" in p for p in paragraphs)
    assert not any("初稿" in p for p in paragraphs)


async def test_docx_no_data_appends_instruction():
    original = _make_docx(["內容"])

    result = await _edit_docx(original, "加上備註段落", None)
    paragraphs = _read_docx_paragraphs(result)

    assert "加上備註段落" in paragraphs


# ── Full handler flow ────────────────────────────────────────────


class MockAdapter:
    def __init__(self, file_bytes: bytes | None):
        self._data = file_bytes
        self.download_calls = 0
        self.fetch_calls = 0

    async def download_media(self, media_id: str) -> bytes | None:
        self.download_calls += 1
        return self._data

    async def fetch_source_file(
        self,
        source: SourceHandleInput,
    ) -> SourceFetchResult | None:
        self.fetch_calls += 1
        if self._data is None:
            return None
        return SourceFetchResult(
            data=self._data,
            filename=source.filename,
            mime_type=source.mime_type,
        )


class MockR2:
    def __init__(self, url: str | None = "https://r2.example.com/test.xlsx"):
        self._url = url

    async def upload(self, data: bytes, content_type: str, extension: str) -> str | None:
        self.last_upload = (data, content_type, extension)
        return self._url


async def test_handler_xlsx_full_flow():
    xlsx_bytes = _make_xlsx([["品名", "數量"], ["水泥漆", 10]])
    adapter = MockAdapter(xlsx_bytes)
    r2 = MockR2("https://r2.example.com/edited.xlsx")

    tool = create_document_edit_tool({"mock": adapter}, r2)
    result = await tool.handler({
        "arguments": {
            "file_ref": "mock:msg_1:report.xlsx",
            "instruction": "追加出貨紀錄",
            "data": json.dumps([["乳膠漆", 5]]),
        }
    })

    assert result["resultType"] == "success"
    assert "edited.xlsx" in result["textResultForLlm"]
    assert "r2.example.com" in result["textResultForLlm"]

    # Verify R2 received valid xlsx
    uploaded_data, ct, ext = r2.last_upload
    assert ext == "xlsx"
    rows = _read_xlsx_rows(uploaded_data)
    assert len(rows) == 3


async def test_handler_docx_full_flow():
    docx_bytes = _make_docx(["第一段"])
    adapter = MockAdapter(docx_bytes)
    r2 = MockR2("https://r2.example.com/edited.docx")

    tool = create_document_edit_tool({"mock": adapter}, r2)
    result = await tool.handler({
        "arguments": {
            "file_ref": "mock:msg_2:memo.docx",
            "instruction": "加上結論",
            "data": json.dumps("結論：專案完成"),
        }
    })

    assert result["resultType"] == "success"
    uploaded_data, _, ext = r2.last_upload
    assert ext == "docx"
    paragraphs = _read_docx_paragraphs(uploaded_data)
    assert "結論：專案完成" in paragraphs


async def test_handler_prefers_center_lookup_when_session_context_present(tmp_path):
    xlsx_bytes = _make_xlsx([["品名", "數量"], ["水泥漆", 10]])
    adapter = MockAdapter(xlsx_bytes)
    store = SqliteFileStore(str(tmp_path / "files.db"))
    await store.initialize()
    center = FileHandleCenter(
        store,
        {"mock": adapter},
        asset_root=tmp_path / "assets",
    )
    handle = await center.register(
        SourceHandleInput(
            route_id="mock:C777",
            platform="mock",
            kind=FileKind.file,
            native_locator="msg_777",
            filename="report.xlsx",
            mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    )
    await center.download_now(handle.file_id)
    adapter.download_calls = 0

    r2 = MockR2("https://r2.example.com/center.xlsx")
    tool = create_document_edit_tool({"mock": adapter}, r2, center)
    result = await tool.handler({
        "session_id": "sid-1",
        "session_context": {
            "sdk_session_id": "sid-1",
            "route_id": "mock:C777",
            "platform": "mock",
            "conversation_id": "C777",
            "chatbot_name": "buddy",
        },
        "arguments": {
            "file_ref": "mock:msg_777:report.xlsx",
            "instruction": "追加出貨紀錄",
            "data": json.dumps([["乳膠漆", 5]]),
        },
    })

    assert result["resultType"] == "success"
    assert adapter.download_calls == 0
    uploaded_data, _, ext = r2.last_upload
    assert ext == "xlsx"
    rows = _read_xlsx_rows(uploaded_data)
    assert len(rows) == 3

    await store.close()


async def test_handler_unsupported_format():
    adapter = MockAdapter(b"%PDF-1.4 fake pdf")
    r2 = MockR2()

    tool = create_document_edit_tool({"mock": adapter}, r2)
    result = await tool.handler({
        "arguments": {
            "file_ref": "mock:msg_3:report.pdf",
            "instruction": "修改",
        }
    })

    assert result["resultType"] == "failure"
    assert "只支援" in result["textResultForLlm"]


async def test_handler_download_failed():
    adapter = MockAdapter(None)
    r2 = MockR2()

    tool = create_document_edit_tool({"mock": adapter}, r2)
    result = await tool.handler({
        "arguments": {
            "file_ref": "mock:msg_4:report.xlsx",
            "instruction": "修改",
        }
    })

    assert result["resultType"] == "failure"
    assert "過期" in result["textResultForLlm"]


async def test_handler_r2_upload_failed():
    xlsx_bytes = _make_xlsx([["A"]])
    adapter = MockAdapter(xlsx_bytes)
    r2 = MockR2(None)  # upload returns None

    tool = create_document_edit_tool({"mock": adapter}, r2)
    result = await tool.handler({
        "arguments": {
            "file_ref": "mock:msg_5:data.xlsx",
            "instruction": "追加",
            "data": json.dumps([["B"]]),
        }
    })

    assert result["resultType"] == "failure"
    assert "上傳失敗" in result["textResultForLlm"]
